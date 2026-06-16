from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedDocumentType(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentType("Only PDF, DOCX, and TXT files are supported")
    if extension == ".pdf":
        return extract_pdf_text(content)
    if extension == ".docx":
        return extract_docx_text(content)
    return content.decode("utf-8", errors="ignore")


def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip())


def extract_docx_text(content: bytes) -> str:
    document = DocxDocument(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    table_cells.append(value)
    return "\n\n".join([*paragraphs, *table_cells])
